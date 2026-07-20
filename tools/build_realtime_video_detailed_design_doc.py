from pathlib import Path

from docx import Document
from docx.enum.section import WD_ORIENT
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor


OUT = Path("/Users/leelay/Documents/robot-mediaserver/实时视频模块详细设计说明书-确认版.docx")


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def set_table_borders(table):
    borders = OxmlElement("w:tblBorders")
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        tag = OxmlElement(f"w:{edge}")
        tag.set(qn("w:val"), "single")
        tag.set(qn("w:sz"), "6")
        tag.set(qn("w:space"), "0")
        tag.set(qn("w:color"), "B8C2CC")
        borders.append(tag)
    table._tbl.tblPr.append(borders)


def set_cell_text(cell, text, bold=False, center=False):
    cell.text = ""
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER if center else WD_ALIGN_PARAGRAPH.LEFT
    run = p.add_run(text)
    run.bold = bold
    run.font.name = "Arial"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "微软雅黑")
    run.font.size = Pt(8.5)
    cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER


def style_document(doc):
    sec = doc.sections[0]
    sec.orientation = WD_ORIENT.LANDSCAPE
    sec.page_width = Cm(29.7)
    sec.page_height = Cm(21.0)
    sec.top_margin = Cm(1.8)
    sec.bottom_margin = Cm(1.6)
    sec.left_margin = Cm(1.8)
    sec.right_margin = Cm(1.8)

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Arial"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "微软雅黑")
    normal.font.size = Pt(10)
    normal.paragraph_format.line_spacing = 1.15
    normal.paragraph_format.space_after = Pt(5)

    title = styles["Title"]
    title.font.name = "Arial"
    title._element.rPr.rFonts.set(qn("w:eastAsia"), "微软雅黑")
    title.font.size = Pt(20)
    title.font.bold = True
    title.font.color.rgb = RGBColor(31, 54, 81)

    for name, size, color in [
        ("Heading 1", 14, RGBColor(31, 78, 121)),
        ("Heading 2", 12, RGBColor(48, 84, 150)),
        ("Heading 3", 10.8, RGBColor(64, 64, 64)),
    ]:
        style = styles[name]
        style.font.name = "Arial"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "微软雅黑")
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = color
        style.paragraph_format.space_before = Pt(10)
        style.paragraph_format.space_after = Pt(5)


def p(doc, text):
    para = doc.add_paragraph()
    run = para.add_run(text)
    run.font.name = "Arial"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "微软雅黑")
    run.font.size = Pt(10)
    return para


def bullet(doc, text):
    para = doc.add_paragraph(style="List Bullet")
    para.paragraph_format.space_after = Pt(2)
    run = para.add_run(text)
    run.font.name = "Arial"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "微软雅黑")
    run.font.size = Pt(9.5)


def number(doc, text):
    para = doc.add_paragraph(style="List Number")
    para.paragraph_format.space_after = Pt(2)
    run = para.add_run(text)
    run.font.name = "Arial"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "微软雅黑")
    run.font.size = Pt(9.5)


def code(doc, lines):
    table = doc.add_table(rows=1, cols=1)
    set_table_borders(table)
    cell = table.cell(0, 0)
    set_cell_shading(cell, "F4F7FA")
    para = cell.paragraphs[0]
    para.paragraph_format.space_before = Pt(4)
    para.paragraph_format.space_after = Pt(4)
    for i, line in enumerate(lines):
        if i:
            para.add_run("\n")
        run = para.add_run(line)
        run.font.name = "Consolas"
        run._element.rPr.rFonts.set(qn("w:eastAsia"), "微软雅黑")
        run.font.size = Pt(8.5)
    doc.add_paragraph()


def table(doc, headers, rows, widths=None):
    t = doc.add_table(rows=1 + len(rows), cols=len(headers))
    set_table_borders(t)
    for c, h in enumerate(headers):
        set_cell_text(t.cell(0, c), h, bold=True, center=True)
        set_cell_shading(t.cell(0, c), "D9EAF7")
    for r, row in enumerate(rows, 1):
        for c, text in enumerate(row):
            set_cell_text(t.cell(r, c), str(text), center=(len(str(text)) < 16))
            if c == 0:
                set_cell_shading(t.cell(r, c), "F2F6FA")
    if widths:
        for row in t.rows:
            for idx, width in enumerate(widths):
                row.cells[idx].width = Cm(width)
    doc.add_paragraph()
    return t


def add_title(doc):
    title = doc.add_paragraph(style="Title")
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.add_run("实时视频模块详细设计说明书")
    sub = doc.add_paragraph()
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = sub.add_run("具身智能装备集成管理平台 / 媒体服务模块 / 可落地开发版")
    run.font.name = "Arial"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "微软雅黑")
    run.font.size = Pt(11)
    run.font.color.rgb = RGBColor(96, 96, 96)


def build():
    doc = Document()
    style_document(doc)
    add_title(doc)

    doc.add_heading("1. 文档目标与设计结论", level=1)
    p(doc, "本文档用于指导实时视频模块的后端、云接入客户端、前端和联调开发。设计基于已确认的方案 A：机器人端统一云接入客户端作为 LiveKit Publisher，平台通过 Media Service 和 LiveKit SFU 提供实时观看能力。")
    table(
        doc,
        ["项目", "内容"],
        [
            ["主方案", "云接入客户端采集双光云台 RTSP 等媒体源，发布为 LiveKit Video Track；平台侧统一管理会话、鉴权、Token、状态和观看。"],
            ["核心原则", "平台对上只暴露统一媒体会话模型，对下由云接入客户端屏蔽 RTSP、厂商 SDK、本地摄像头等差异。"],
            ["实时视频范围", "创建会话、启动发布、前端观看、状态推送、通道切换、停止观看、异常处理、资源释放。"],
            ["关联但不展开", "录制、抓拍、回放只定义衔接边界：录制主路径 LiveKit Egress；抓拍采用前端即时预览 + 服务端基于 LiveKit Track 权威入库；回放主路径播放平台已存储文件。"],
        ],
        [3.2, 20],
    )

    doc.add_heading("2. 术语与角色", level=1)
    table(
        doc,
        ["术语", "说明"],
        [
            ["云接入客户端", "部署在机器人端的统一抽象层，负责设备协议适配、媒体采集、LiveKit 发布、状态上报和控制指令处理。"],
            ["Media Service", "平台媒体服务，负责实时视频会话、权限、Token、LiveKit 元数据、状态推送和资源释放策略。"],
            ["LiveKit", "实时媒体 SFU，负责 Room、Participant、Track、订阅、发布和媒体转发。"],
            ["Video Session", "平台侧业务会话，描述一次实时观看链路，不等同于 LiveKit Room，但会关联 Room 和 Track。"],
            ["Channel", "双光云台媒体通道，如 visible、thermal、fusion。"],
            ["Track", "LiveKit 中的视频或音频媒体轨道。实时视频模块主要管理视频 Track。"],
        ],
        [3, 20],
    )

    doc.add_heading("3. 总体架构", level=1)
    code(
        doc,
        [
            "双光云台 RTSP / 本地摄像头 / 厂商 SDK",
            "        |",
            "机器人端云接入客户端",
            "  - RTSP Adapter / Device Adapter",
            "  - LiveKit Publisher",
            "  - Command Handler",
            "  - Status Reporter",
            "        |  WebRTC/SRTP 发布视频 Track",
            "LiveKit SFU",
            "        |  WebRTC/SRTP 订阅视频 Track",
            "Web 管理端 / 手机 App",
            "",
            "控制指令：平台 -> Robot Control Service/MQTT -> 云接入客户端",
            "业务状态：Media Service -> WebSocket -> Web/App",
            "业务接口：Web/App -> API Gateway -> Media Service",
        ],
    )

    doc.add_heading("4. 模块职责边界", level=1)
    table(
        doc,
        ["模块", "负责", "不负责"],
        [
            ["Media Service", "视频会话创建、权限校验、Token 生成、Room/Track 元数据、状态机、WebSocket 状态、资源释放、错误码。", "不直接拉取 RTSP；不控制云台 PTZ；不传输大媒体流。"],
            ["云接入客户端", "拉取/采集媒体源，发布 LiveKit Track，接收开始/停止/切换通道指令，上报 ACK、发布成功、断流、重连等状态。", "不做平台用户权限决策；不分发前端观看 Token；不维护平台媒体资产索引。"],
            ["Robot Control Service/MQTT", "设备在线状态、媒体控制命令、命令回执、异常事件、客户端心跳。", "不承载视频流；不替代 WebSocket 给前端推业务状态。"],
            ["WebSocket", "向前端推送 video.session.* 业务状态。", "不下发底层设备控制命令；不传输视频流。"],
            ["LiveKit", "媒体房间、发布订阅、SFU 转发、断连事件。", "不保存业务会话，不替代 Media Service 的权限和状态机。"],
            ["Robot Control Service", "云台 PTZ、变焦、预置位、补光灯、雨刷、机械臂等动作。", "Media Service 只处理媒体通道，不处理设备机械动作。"],
        ],
        [3, 11, 9],
    )

    doc.add_heading("5. 核心业务流程", level=1)
    doc.add_heading("5.1 创建实时视频会话", level=2)
    code(
        doc,
        [
            "1. 前端调用 POST /internal/media/video-sessions，提交 robotId、deviceId、channel、quality。",
            "2. Media Service 校验用户权限、机器人在线状态、设备媒体能力和并发资源配额。",
            "3. Media Service 创建或复用 LiveKit Room，生成机器人端发布 Token。",
            "4. Media Service 通过 MQTT 下发 media.video.start 指令给云接入客户端。",
            "5. 云接入客户端返回 ACK，并拉取双光云台 RTSP 或本地媒体源。",
            "6. 云接入客户端以 Publisher 身份加入 LiveKit Room，发布 Video Track。",
            "7. Media Service 监听到 Track published，更新会话状态为 STREAMING。",
            "8. Media Service 生成前端观看 Token，返回给前端或由前端另行获取。",
            "9. 前端加入 LiveKit Room，订阅目标 Track，WebSocket 接收状态变化。",
        ],
    )
    doc.add_heading("5.2 停止观看与资源释放", level=2)
    code(
        doc,
        [
            "1. 前端调用 POST /internal/media/video-sessions/{sessionId}/stop。",
            "2. Media Service 将当前用户从观看者列表移除，更新 viewerCount。",
            "3. 若 viewerCount > 0，仅推送 viewer.changed，不停止客户端发布。",
            "4. 若 viewerCount = 0，进入 IDLE_WAIT，默认等待 60 秒。",
            "5. 等待期内有新观看者加入，则恢复 STREAMING 并复用原 Track。",
            "6. 等待期结束仍无人观看，Media Service 下发 media.video.stop。",
            "7. 云接入客户端停止发布 Track，释放 RTSP/本地媒体资源。",
            "8. Media Service 关闭会话或标记 CLOSED，推送 video.session.closed。",
        ],
    )
    doc.add_heading("5.3 通道切换", level=2)
    code(
        doc,
        [
            "1. 前端调用 POST /internal/media/video-sessions/{sessionId}/switch-channel。",
            "2. Media Service 校验目标通道 visible/thermal/fusion 是否被设备支持。",
            "3. Media Service 下发 media.video.switchChannel 指令。",
            "4. 云接入客户端切换 RTSP 源或设备采集通道。",
            "5. 若 Track 可平滑替换，则保持 Room 不变并发布新 Track；否则先中断旧 Track 再发布新 Track。",
            "6. Media Service 先推送 video.track.switching；新 Track 可用后推送 video.session.streaming 或 video.track.published。",
        ],
    )
    doc.add_heading("5.4 异常恢复", level=2)
    code(
        doc,
        [
            "客户端 ACK 超时：REQUESTING_CLIENT -> TIMEOUT，释放 Room，推送失败。",
            "Track 发布超时：CLIENT_ACKED -> FAILED，下发 stop，记录 LK_PUBLISH_TIMEOUT。",
            "Track 中断：STREAMING -> INTERRUPTED，允许客户端重连，超过阈值转 FAILED。",
            "客户端重连：若 session 未过期且 robotId/deviceId/channel 匹配，可复用 Room 并更新 Track。",
            "前端异常退出：依赖 LiveKit participant disconnected 和心跳修正 viewerCount。",
        ],
    )

    doc.add_heading("6. 状态机设计", level=1)
    table(
        doc,
        ["状态", "含义", "进入条件", "退出条件/下一状态", "是否推送前端"],
        [
            ["INIT", "会话已创建，尚未下发客户端指令。", "API 创建会话成功。", "下发 start 成功 -> REQUESTING_CLIENT；失败 -> FAILED。", "是"],
            ["REQUESTING_CLIENT", "等待云接入客户端 ACK。", "MQTT start 指令已发送。", "收到 ACK -> CLIENT_ACKED；超时 -> TIMEOUT。", "是"],
            ["CLIENT_ACKED", "客户端确认开始准备发布媒体。", "收到 start ACK。", "Room/Token 就绪 -> ROOM_READY；发布超时 -> FAILED。", "是"],
            ["ROOM_READY", "LiveKit Room 已准备，等待 Track 发布或前端加入。", "Room 创建/复用成功。", "Track published -> STREAMING；超时 -> FAILED。", "是"],
            ["STREAMING", "视频 Track 正常发布，可观看。", "LiveKit Track published。", "停止 -> STOPPING；断流 -> INTERRUPTED。", "是"],
            ["INTERRUPTED", "Track 中断或客户端重连中。", "LiveKit 断开、客户端上报断流。", "恢复 -> STREAMING；超时 -> FAILED。", "是"],
            ["IDLE_WAIT", "无人观看，延迟释放。", "viewerCount 降为 0。", "新观看者加入 -> STREAMING；超时 -> STOPPING。", "可选"],
            ["STOPPING", "正在停止发布和释放资源。", "用户停止或空闲释放。", "客户端确认停止 -> CLOSED；超时 -> CLOSED_WITH_WARN。", "是"],
            ["CLOSED", "会话正常关闭。", "客户端停止成功或业务释放。", "终态。", "是"],
            ["TIMEOUT", "等待客户端或 Track 超时。", "ACK/发布/重连超过阈值。", "终态或用户重试新建。", "是"],
            ["FAILED", "会话失败。", "权限、资源、LiveKit、客户端异常。", "终态或用户重试新建。", "是"],
        ],
        [2.3, 5.3, 5.2, 5.6, 2.2],
    )
    table(
        doc,
        ["参数", "建议值", "说明"],
        [
            ["clientAckTimeout", "10s", "客户端未 ACK 则会话超时。"],
            ["trackPublishTimeout", "20s", "客户端 ACK 后未发布 Track 则失败。"],
            ["interruptedGrace", "15s", "Track 短暂中断允许恢复。"],
            ["idleReleaseDelay", "60s", "最后观看者退出后延迟释放。"],
            ["tokenTtl", "5-10min", "发布/观看 Token 短时有效，按需刷新。"],
        ],
        [4, 4, 15],
    )

    doc.add_heading("7. API 设计", level=1)
    table(
        doc,
        ["接口", "方法", "说明", "主要状态码"],
        [
            ["/internal/media/video-sessions", "POST", "创建实时视频会话。", "200/201、400、403、409、423、503"],
            ["/internal/media/video-sessions/{sessionId}", "GET", "查询会话详情。", "200、403、404"],
            ["/internal/media/video-sessions/{sessionId}/token", "POST", "获取前端观看 Token。", "200、403、404、409"],
            ["/internal/media/video-sessions/{sessionId}/switch-channel", "POST", "切换 visible/thermal/fusion。", "200、400、403、404、409"],
            ["/internal/media/video-sessions/{sessionId}/stop", "POST", "停止当前用户观看或释放会话。", "200、403、404"],
            ["/internal/media/video-sessions/{sessionId}/snapshots", "POST", "从当前 LiveKit Track 截帧，生成抓拍图片。", "202、403、404、409、503"],
        ],
        [6, 2, 11, 4],
    )
    doc.add_heading("7.1 创建会话请求/响应", level=2)
    code(
        doc,
        [
            "POST /internal/media/video-sessions",
            "Request:",
            "{",
            '  "robotId": "test111",',
            '  "deviceId": "gimbal-001",',
            '  "channel": "visible",',
            '  "quality": "sub",',
            '  "reuse": true,',
            '  "clientRequestId": "uuid"',
            "}",
            "",
            "Response:",
            "{",
            '  "sessionId": "vs_20260518_000001",',
            '  "status": "REQUESTING_CLIENT",',
            '  "roomName": "media.test111.gimbal-001.visible",',
            '  "viewerToken": "jwt",',
            '  "livekitUrl": "wss://livekit.example.com",',
            '  "expiresAt": "2026-05-18T10:30:00+08:00"',
            "}",
        ],
    )
    doc.add_heading("7.2 抓拍请求/响应", level=2)
    code(
        doc,
        [
            "POST /internal/media/video-sessions/{sessionId}/snapshots",
            "Request:",
            "{",
            '  "trackSid": "TR_xxx",',
            '  "reason": "manual_abnormal",',
            '  "clientCapturedAt": "2026-05-18T10:21:31.245+08:00",',
            '  "clientPreviewObjectKey": "snapshots/preview/snap_xxx.jpg",',
            '  "previewImageHash": "sha256:optional",',
            '  "remark": "发现异常画面"',
            "}",
            "",
            "Response:",
            "{",
            '  "snapshotId": "snap_20260518_000001",',
            '  "status": "PROCESSING",',
            '  "mode": "livekit_track",',
            '  "previewAccepted": true,',
            '  "createdAt": "2026-05-18T10:21:31+08:00"',
            "}",
        ],
    )

    doc.add_heading("7.3 抓拍详细设计", level=2)
    p(doc, "抓拍的业务语义是操作员正在观看实时视频时发现异常，点击抓拍后尽量保存接近点击瞬间的画面。因此抓拍不采用纯前端截图作为正式证据，也不只依赖点击后服务端临时订阅 Track，而是采用前端即时预览与服务端权威入库结合的设计。")
    table(
        doc,
        ["设计项", "规则"],
        [
            ["主路径", "前端点击抓拍后立即从当前 video 元素截取预览图，同时调用后端抓拍接口；服务端由 Snapshot Worker 基于 LiveKit Track 生成正式抓拍图片并入库。"],
            ["即时预览", "前端预览用于提升操作反馈速度，可上传为 previewObjectKey，但不作为最终权威图片。"],
            ["权威图片", "服务端生成 officialObjectKey，写入 MinIO 和 media_snapshot，作为审计、检索和业务流转使用的正式图片。"],
            ["第一版落地", "Snapshot Worker 可在收到抓拍任务后加入 LiveKit Room，订阅目标 Track，等待最近一帧并编码上传。"],
            ["增强版落地", "会话进入 STREAMING 后，Snapshot Worker 订阅活跃 Track，维护最近 1-3 秒帧缓存；抓拍时取最接近 clientCapturedAt/serverReceivedAt 的帧。"],
            ["超时策略", "3-5 秒内未获取到有效视频帧则失败，错误码 SNAPSHOT_TIMEOUT；Track 中断则返回 TRACK_INTERRUPTED。"],
            ["审计字段", "记录 sessionId、trackSid、robotId、deviceId、channel、createdBy、clientCapturedAt、officialCapturedAt、source、timeDeltaMs。"],
        ],
        [4, 19],
    )
    code(
        doc,
        [
            "抓拍流程：",
            "1. 用户在实时视频画面中发现异常并点击抓拍。",
            "2. 前端立即 drawImage(video) 生成预览图，界面展示抓拍反馈。",
            "3. 前端调用 snapshots 接口，提交 sessionId、trackSid、clientCapturedAt、reason、remark。",
            "4. Media Service 校验会话处于 STREAMING、用户有权限、Track 存在。",
            "5. Media Service 创建 media_snapshot 记录，状态 PROCESSING。",
            "6. Snapshot Worker 订阅目标 LiveKit Track，或从最近帧缓存取最接近点击时间的帧。",
            "7. Snapshot Worker 编码 JPEG/PNG，上传 MinIO，回写 officialObjectKey。",
            "8. Media Service 推送 snapshot.completed 或 snapshot.failed。",
        ],
    )

    doc.add_heading("8. WebSocket 事件设计", level=1)
    table(
        doc,
        ["事件", "触发时机", "关键字段"],
        [
            ["robot.state", "机器人客户端在线状态、离线状态或控制状态变化。", "robotId、clientId、name、type、battery、status、onlineStatus、controlMode、stateSeq、cameras、devices、timestamp"],
            ["video.session.created / video.session.reused", "会话创建或复用。", "VideoSessionResponse"],
            ["video.room.ready", "Room 创建完成或客户端上报 publishing。", "sessionId、roomName，或 VideoSessionResponse"],
            ["video.client.requested", "已请求客户端启动推流。", "sessionId、commandId、timeoutSeconds"],
            ["video.track.switching", "切换通道或清晰度。", "sessionId、commandId"],
            ["video.track.published", "LiveKit Track 发布成功。", "VideoSessionResponse，包含 trackSid、trackName"],
            ["video.session.streaming", "进入可观看状态。", "VideoSessionResponse"],
            ["video.viewer.changed", "观看人数变化。", "VideoSessionResponse"],
            ["video.session.interrupted", "Track 中断或客户端重连。", "sessionId、message"],
            ["video.session.failed", "会话失败。", "sessionId、errorCode、message"],
            ["video.session.idle_wait", "最后一个观看者离开，等待释放。", "sessionId、idleReleaseDelaySeconds"],
            ["video.session.stopping / video.session.closed", "会话释放或关闭。", "VideoSessionResponse"],
            ["video.session.restart / video.session.auto_restart / video.client.online_restart", "手动、自动或上线恢复重启。", "sessionId、commandId"],
            ["video.intercom.starting / active / stopping / closed", "对讲状态变化。", "VideoSessionResponse，active 时包含 robotAudioTrackSid、robotAudioTrackName"],
            ["video.intercom.interrupted / failed / status", "对讲异常或未显式映射状态。", "sessionId、message；失败时包含 errorCode"],
            ["snapshot.requested", "创建抓拍任务。", "snapshotId、sessionId、trackSid、source、createdBy"],
            ["snapshot.completed", "抓拍正式图片生成成功。", "snapshotId、officialObjectKey、capturedAt、source"],
            ["snapshot.failed", "抓拍失败。", "snapshotId、errorCode、message"],
        ],
        [5, 8, 10],
    )
    code(
        doc,
        [
            "WebSocket Payload 示例：",
            "{",
            '  "event": "video.track.published",',
            '  "timestamp": "2026-05-18T10:20:11+08:00",',
            '  "data": {',
            '    "sessionId": "vs_20260518_000001",',
            '    "roomName": "media.test111.gimbal-001.visible",',
            '    "trackSid": "TR_xxx",',
            '    "trackName": "video.visible.sub",',
            '    "channel": "visible"',
            "  }",
            "}",
        ],
    )

    doc.add_heading("9. MQTT Topic 与 Payload", level=1)
    table(
        doc,
        ["Topic", "方向", "用途"],
        [
            ["robot/{robotId}/media/video/start", "平台 -> 客户端", "启动指定设备和通道的视频发布。"],
            ["robot/{robotId}/media/video/stop", "平台 -> 客户端", "停止指定会话的视频发布。"],
            ["robot/{robotId}/media/video/switch-channel", "平台 -> 客户端", "切换双光通道。"],
            ["robot/{robotId}/media/video/ack", "客户端 -> 平台", "命令 ACK。"],
            ["robot/{robotId}/media/video/status", "客户端 -> 平台", "发布状态、断流、恢复、错误上报。"],
            ["robot/{robotId}/heartbeat", "客户端 -> 平台", "机器人和云接入客户端在线状态。"],
        ],
        [8, 4, 12],
    )
    code(
        doc,
        [
            "media.video.start Payload:",
            "{",
            '  "commandId": "cmd_xxx",',
            '  "sessionId": "vs_20260518_000001",',
            '  "robotId": "test111",',
            '  "deviceId": "gimbal-001",',
            '  "channel": "visible",',
            '  "quality": "sub",',
            '  "livekitUrl": "wss://livekit.example.com",',
            '  "roomName": "media.test111.gimbal-001.visible",',
            '  "publisherToken": "jwt",',
            '  "publishIdentity": "robot:test111:gimbal-001",',
            '  "expiresAt": "2026-05-18T10:30:00+08:00"',
            "}",
        ],
    )

    doc.add_heading("10. LiveKit 命名与权限", level=1)
    table(
        doc,
        ["对象", "规则", "示例"],
        [
            ["Room", "media.{robotId}.{deviceId}.{channel}", "media.test111.gimbal-001.visible"],
            ["机器人 Participant", "robot:{robotId}:{deviceId}", "robot:test111:gimbal-001"],
            ["前端 Participant", "user:{userId}:{clientType}", "user:u1001:web"],
            ["Track Name", "video.{channel}.{quality}", "video.visible.sub"],
            ["发布 Token", "只允许加入指定 Room 并发布指定 Track，不允许订阅其他机器人。", "canPublish=true, canSubscribe=false"],
            ["观看 Token", "只允许加入授权 Room 并订阅 Track，不允许发布。", "canPublish=false, canSubscribe=true"],
        ],
        [4, 12, 7],
    )
    bullet(doc, "Room/Track 命名不得包含密码、RTSP 地址、项目密级、现场位置等敏感信息。")
    bullet(doc, "Token 短时有效，Media Service 按会话生成，不复用长期凭据。")
    bullet(doc, "机器人端发布 Token 和前端观看 Token 分离，权限最小化。")

    doc.add_heading("11. 数据模型设计", level=1)
    table(
        doc,
        ["表", "用途", "关键字段"],
        [
            ["media_video_session", "实时视频业务会话。", "session_id、robot_id、device_id、channel、quality、room_name、status、viewer_count、created_by、started_at、ended_at、last_error_code"],
            ["media_track", "LiveKit Track 元数据。", "track_id、session_id、track_sid、track_name、participant_identity、kind、channel、quality、published_at、unpublished_at"],
            ["media_session_viewer", "观看者记录。", "id、session_id、user_id、participant_identity、joined_at、left_at、client_type"],
            ["media_event_log", "状态和异常日志。", "id、session_id、event_type、event_payload、trace_id、created_at"],
            ["media_snapshot", "抓拍记录。", "snapshot_id、session_id、track_sid、robot_id、device_id、channel、official_object_key、preview_object_key、source、time_delta_ms、reason、created_by、created_at"],
        ],
        [4, 7, 13],
    )
    doc.add_heading("11.1 索引建议", level=2)
    bullet(doc, "media_video_session：robot_id + device_id + channel + status，用于复用会话和判断占用。")
    bullet(doc, "media_video_session：created_by + created_at，用于用户历史查询。")
    bullet(doc, "media_track：session_id + track_sid，用于 Track 事件回写。")
    bullet(doc, "media_event_log：session_id + created_at，用于联调排错。")
    bullet(doc, "media_snapshot：robot_id + device_id + created_at，用于媒体检索。")

    doc.add_heading("12. 错误码设计", level=1)
    table(
        doc,
        ["错误码", "HTTP", "含义", "处理建议"],
        [
            ["ROBOT_OFFLINE", "503", "机器人或云接入客户端离线。", "提示设备离线，禁止创建会话。"],
            ["DEVICE_NOT_FOUND", "404", "设备或云台不存在。", "检查 robotId/deviceId。"],
            ["CHANNEL_UNSUPPORTED", "400", "设备不支持目标通道。", "前端刷新设备能力。"],
            ["MEDIA_PERMISSION_DENIED", "403", "用户无观看权限。", "走权限申请或切换账号。"],
            ["MEDIA_RESOURCE_LIMIT", "423", "并发资源达到上限。", "提示稍后重试或释放其他会话。"],
            ["CLIENT_ACK_TIMEOUT", "504", "客户端未在规定时间 ACK。", "检查 MQTT、客户端状态。"],
            ["LK_ROOM_CREATE_FAILED", "502", "LiveKit Room 创建失败。", "检查 LiveKit 服务。"],
            ["LK_PUBLISH_TIMEOUT", "504", "客户端未按时发布 Track。", "检查媒体源和客户端日志。"],
            ["TRACK_INTERRUPTED", "409", "Track 当前中断。", "前端展示恢复中或重试。"],
            ["TOKEN_EXPIRED", "401", "Token 过期。", "重新申请 Token。"],
            ["SNAPSHOT_FAILED", "503", "抓拍失败。", "提示重试，记录 Track 状态。"],
            ["SNAPSHOT_TIMEOUT", "504", "服务端在规定时间内未获取到有效视频帧。", "提示重试，检查 Track 是否正常和 Snapshot Worker 状态。"],
        ],
        [4, 2, 8, 10],
    )

    doc.add_heading("13. 并发与资源策略", level=1)
    table(
        doc,
        ["策略", "规则"],
        [
            ["多人观看", "同一 robotId + deviceId + channel + quality 已有 STREAMING 会话时，默认复用 Room/Track，不重复启动客户端发布。"],
            ["视频墙", "视频墙最多按 16 路设计，优先订阅 sub 码流或低清 Track；单路详情切换 main 码流。"],
            ["高清限制", "单路详情最高支持 2K，平台配置 maxHdSessions，超过后拒绝或降级为 sub。"],
            ["空闲释放", "viewerCount 为 0 后进入 IDLE_WAIT，默认 60 秒后释放客户端媒体资源。"],
            ["对讲占用", "同一机器人同一时间只允许一个主控对讲；实时视频可多人观看。"],
            ["重试策略", "客户端发布失败不无限重试；同一会话最多重试 2 次，超过进入 FAILED。"],
        ],
        [5, 18],
    )

    doc.add_heading("14. 安全设计", level=1)
    bullet(doc, "API 访问依赖平台 Auth Service，Media Service 只接受已认证用户请求。")
    bullet(doc, "发布 Token、观看 Token、API Token 三类凭据分离。")
    bullet(doc, "机器人端只能发布与自身 robotId/deviceId 匹配的 Track。")
    bullet(doc, "前端只能订阅授权范围内机器人和通道。")
    bullet(doc, "MQTT 指令包含 commandId、sessionId、expiresAt 和签名/摘要字段，客户端拒绝过期或重复命令。")
    bullet(doc, "日志中不得打印完整 Token、RTSP 密码、设备密钥。")

    doc.add_heading("15. 开发模块划分", level=1)
    table(
        doc,
        ["端", "模块", "职责"],
        [
            ["后端", "VideoSessionService", "会话创建、复用、停止、状态机推进。"],
            ["后端", "LiveKitService", "Room、Token、Track 事件处理。"],
            ["后端", "RobotMediaCommandService", "MQTT 指令封装、发送、ACK 关联。"],
            ["后端", "MediaStatusPublisher", "WebSocket 事件发布。"],
            ["后端", "SnapshotService", "订阅/抓取 LiveKit 当前 Track 帧并写入 MinIO。"],
            ["后端", "MediaSessionRepository", "会话、Track、观看者、事件日志持久化。"],
            ["客户端", "RtspAdapter", "连接双光云台 RTSP，处理 visible/thermal/fusion。"],
            ["客户端", "LiveKitPublisher", "加入 Room，发布 Video Track，断线重连。"],
            ["客户端", "CommandHandler", "处理 start、stop、switch-channel 指令。"],
            ["客户端", "StatusReporter", "上报 ACK、发布成功、断流、恢复和错误。"],
            ["前端", "VideoPlayer", "申请会话/Token，加入 Room，展示视频和状态。"],
        ],
        [3, 5, 15],
    )

    doc.add_heading("16. 联调与验收用例", level=1)
    table(
        doc,
        ["用例", "步骤", "预期结果"],
        [
            ["正常观看", "机器人在线，创建 visible 会话，前端加入 Room。", "10-30 秒内进入 STREAMING，画面可播放。"],
            ["多人观看复用", "两个用户观看同一机器人同一通道。", "复用同一 Room/Track，viewerCount=2，不重复下发 start。"],
            ["通道切换", "从 visible 切换到 thermal。", "推送 track.switched，前端画面切到热成像。"],
            ["最后观看者退出", "所有前端停止观看。", "进入 IDLE_WAIT，60 秒后下发 stop 并 CLOSED。"],
            ["机器人离线", "离线机器人创建会话。", "返回 ROBOT_OFFLINE，不创建 LiveKit 发布任务。"],
            ["ACK 超时", "客户端不返回 ACK。", "状态 TIMEOUT，推送 client ack timeout。"],
            ["Track 中断恢复", "客户端短暂断网后恢复。", "进入 INTERRUPTED 后恢复 STREAMING。"],
            ["权限不足", "无权限用户请求观看。", "返回 MEDIA_PERMISSION_DENIED。"],
            ["Token 过期", "使用过期 Token 加入 Room。", "加入失败，前端重新申请 Token。"],
            ["观看抓拍", "STREAMING 状态下调用 snapshots。", "生成 snapshot 记录，图片写入 MinIO。"],
        ],
        [4, 10, 10],
    )

    doc.add_heading("17. 开发顺序建议", level=1)
    number(doc, "实现 Media Service 会话表、状态机、创建/查询/停止接口。")
    number(doc, "接入 LiveKit Room 和 Token 生成，打通前端加入空 Room。")
    number(doc, "实现 MQTT start/stop 指令和客户端 ACK 回调。")
    number(doc, "实现云接入客户端 RTSP Adapter + LiveKit Publisher，发布 visible 通道。")
    number(doc, "实现 Track published 事件回写和 WebSocket 状态推送。")
    number(doc, "实现多人观看复用、空闲释放、重复请求处理。")
    number(doc, "实现 thermal 通道和 switch-channel。")
    number(doc, "实现异常处理、错误码、事件日志和联调用例。")
    number(doc, "实现抓拍两段式能力：前端即时预览，服务端 Snapshot Worker 订阅 LiveKit Track 生成正式图片。")
    number(doc, "根据准确性要求增强 Snapshot Worker 最近帧缓存，保留活跃 Track 最近 1-3 秒抽样帧。")
    number(doc, "补充录制 Egress 和回放模块的衔接接口。")

    doc.add_heading("18. 已确认设计输入", level=1)
    table(
        doc,
        ["确认项", "确认结果", "设计影响"],
        [
            [
                "云接入客户端运行环境",
                "统一客户端部署在机器人侧，主控板为 Jetson Orin NX 16GB，含 256GB SSD。",
                "具备运行 LiveKit SDK、GStreamer/FFmpeg、MQTT 客户端、状态上报服务和本地日志缓存的硬件条件；媒体处理优先考虑 GStreamer，FFmpeg 作为调试和兜底工具。",
            ],
            [
                "双光云台 RTSP 地址管理",
                "RTSP 地址固定配置在机器人侧云接入客户端；平台 start 指令不下发完整 RTSP URL。",
                "平台只下发 robotId、deviceId、channel、quality、sessionId、roomName 和 publisherToken；云接入客户端根据本地配置解析 visible/thermal 的 main/sub RTSP 源。",
            ],
            [
                "LiveKit 部署方式",
                "LiveKit 采用平台自部署；第一阶段以内网环境跑通为目标，暂不强依赖公网 TURN/STUN。",
                "Nginx 作为统一 HTTPS/WSS 入口代理前端、API、业务 WebSocket 和 LiveKit 信令；WebRTC 媒体端口按 LiveKit 内网部署要求暴露，后续公网化再补 TURN/TLS。",
            ],
            [
                "访问路径",
                "机器人端通过边缘网关访问平台 LiveKit、API/MQTT 服务；前端采用前后端分离，静态资源部署在 Nginx。",
                "机器人端保持主动出站访问，前端通过 HTTPS/WSS 访问平台服务，符合当前安全和部署规划。",
            ],
            [
                "并发容量目标",
                "一期最少支持 2 台机器人同时接入实时视频；设计按 16 路视频墙能力预留。",
                "视频墙默认使用 sub 码流/低清 Track；并发上限通过配置控制，不在代码中写死。",
            ],
            [
                "清晰度目标",
                "单路详情最高支持 2K 主码流；视频墙使用低清/子码流。",
                "建议视频墙 480p/720p、5-10fps；单路详情 1080p/2K、15-25fps，具体按设备和网络能力自适应。",
            ],
            [
                "多人观看",
                "必须支持多人同时观看同一路视频。",
                "同一 robotId + deviceId + channel + quality 已处于 STREAMING 时，后续观看者复用同一 Room/Track，不重复下发 start，不重复发布 Track。",
            ],
            [
                "抓拍落地方式",
                "抓拍基于正在打开观看的视频；采用前端即时预览 + 服务端订阅 LiveKit Track 权威入库。",
                "第一版可点击后按需订阅 Track；增强版 Snapshot Worker 维护最近 1-3 秒帧缓存，快速取最接近点击时间的异常画面帧。",
            ],
        ],
        [4, 9, 14],
    )

    doc.add_heading("19. 后续待细化事项", level=1)
    bullet(doc, "Jetson 侧具体软件环境仍需确认：Ubuntu/JetPack 版本、LiveKit SDK 语言栈、GStreamer 插件可用性。")
    bullet(doc, "双光云台 RTSP 具体路径、认证方式、主/子码流参数、可见光/热成像切换机制需在设备到场后实测。")
    bullet(doc, "LiveKit 内网部署的 UDP/TCP 端口、Nginx 反向代理配置、证书和域名需要由部署方案进一步固化。")
    bullet(doc, "16 路视频墙和 2K 单路详情需要通过 LiveKit、边缘网关、浏览器解码和机器人端上行带宽压测校准最终配置。")
    bullet(doc, "Snapshot Worker 的具体实现选型需要结合后端语言栈决定：直接订阅 LiveKit Track、独立媒体 Worker，或与 Egress 能力组合。")

    doc.save(OUT)
    return OUT


if __name__ == "__main__":
    print(build())
