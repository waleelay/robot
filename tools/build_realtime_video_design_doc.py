from pathlib import Path

from docx import Document
from docx.enum.section import WD_ORIENT
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor


OUT = Path("/Users/leelay/Documents/robot-mediaserver/实时视频模块设计方案-方案A选择说明-含边界职责.docx")


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def set_cell_text(cell, text, bold=False):
    cell.text = ""
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER if "\n" not in text and len(text) < 18 else WD_ALIGN_PARAGRAPH.LEFT
    run = p.add_run(text)
    run.bold = bold
    run.font.name = "Arial"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "微软雅黑")
    run.font.size = Pt(9)
    cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER


def set_table_borders(table):
    tbl = table._tbl
    tbl_pr = tbl.tblPr
    borders = OxmlElement("w:tblBorders")
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        tag = OxmlElement(f"w:{edge}")
        tag.set(qn("w:val"), "single")
        tag.set(qn("w:sz"), "6")
        tag.set(qn("w:space"), "0")
        tag.set(qn("w:color"), "B8C2CC")
        borders.append(tag)
    tbl_pr.append(borders)


def add_bullet(doc, text):
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.space_after = Pt(3)
    run = p.add_run(text)
    run.font.name = "Arial"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "微软雅黑")
    run.font.size = Pt(10.5)


def add_number(doc, text):
    p = doc.add_paragraph(style="List Number")
    p.paragraph_format.space_after = Pt(3)
    run = p.add_run(text)
    run.font.name = "Arial"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "微软雅黑")
    run.font.size = Pt(10.5)


def add_code_block(doc, lines):
    table = doc.add_table(rows=1, cols=1)
    set_table_borders(table)
    cell = table.cell(0, 0)
    set_cell_shading(cell, "F4F7FA")
    p = cell.paragraphs[0]
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(4)
    for i, line in enumerate(lines):
        if i:
            p.add_run("\n")
        run = p.add_run(line)
        run.font.name = "Consolas"
        run._element.rPr.rFonts.set(qn("w:eastAsia"), "微软雅黑")
        run.font.size = Pt(9)
    doc.add_paragraph()


def style_document(doc):
    sec = doc.sections[0]
    sec.orientation = WD_ORIENT.LANDSCAPE
    sec.page_width = Cm(29.7)
    sec.page_height = Cm(21.0)
    sec.top_margin = Cm(2.2)
    sec.bottom_margin = Cm(2.0)
    sec.left_margin = Cm(2.2)
    sec.right_margin = Cm(2.2)

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Arial"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "微软雅黑")
    normal.font.size = Pt(10.5)
    normal.paragraph_format.line_spacing = 1.18
    normal.paragraph_format.space_after = Pt(6)

    title = styles["Title"]
    title.font.name = "Arial"
    title._element.rPr.rFonts.set(qn("w:eastAsia"), "微软雅黑")
    title.font.size = Pt(20)
    title.font.bold = True
    title.font.color.rgb = RGBColor(31, 54, 81)

    for name, size, color in [
        ("Heading 1", 15, RGBColor(31, 78, 121)),
        ("Heading 2", 12.5, RGBColor(48, 84, 150)),
        ("Heading 3", 11.5, RGBColor(64, 64, 64)),
    ]:
        style = styles[name]
        style.font.name = "Arial"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "微软雅黑")
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = color
        style.paragraph_format.space_before = Pt(12)
        style.paragraph_format.space_after = Pt(6)


def add_kv_table(doc, rows):
    table = doc.add_table(rows=len(rows), cols=2)
    set_table_borders(table)
    table.columns[0].width = Cm(4)
    table.columns[1].width = Cm(12)
    for i, (k, v) in enumerate(rows):
        set_cell_text(table.cell(i, 0), k, bold=True)
        set_cell_shading(table.cell(i, 0), "EAF2F8")
        set_cell_text(table.cell(i, 1), v)
    doc.add_paragraph()


def add_comparison_table(doc):
    headers = ["维度", "方案 A：云接入客户端作为 LiveKit Publisher", "方案 B：平台侧拉取 RTSP/RTMP/GB28181"]
    rows = [
        ["核心思路", "机器人端统一客户端屏蔽设备差异，向平台发布标准音视频 Track。", "机器人或设备提供标准流地址，平台媒体服务集中拉流、转码或转推。"],
        ["与当前架构匹配度", "高。原架构已有云接入客户端、WebRTC(SRTP)、SFU、接入认证与通信保护。", "中。适合作为兼容能力，但会弱化云接入客户端的媒体抽象作用。"],
        ["多机器人接入", "平台只面对统一媒体协议，新增机器人主要扩展客户端适配器。", "平台要管理多协议、多网络、多拉流任务，接入复杂度集中到云侧。"],
        ["实时视频", "低延迟链路更直接，前端统一订阅 LiveKit Room/Track。", "可快速兼容 RTSP，但需要转码/转推，延迟与资源消耗更高。"],
        ["语音对讲", "天然适合。音频 Track、房间、权限和会话控制可复用。", "不理想。RTSP 类协议不适合双向语音，需要额外通道补齐。"],
        ["网络适应性", "机器人主动出站连接平台，适合 4G/5G、NAT 和移动网络。", "平台访问机器人内网 RTSP 往往受网络、端口和安全策略限制。"],
        ["平台运维压力", "媒体转接压力分散到机器人端/边缘端，平台侧以 SFU 转发为主。", "平台侧容易成为 FFmpeg/Ingress 拉流中心，资源和故障排查压力更大。"],
        ["设备兼容", "客户端内部仍需支持 RTSP、厂商 SDK、本地摄像头、麦克风等适配。", "兼容存量视频设备更直接，是必要的兜底方式。"],
    ]
    table = doc.add_table(rows=1 + len(rows), cols=3)
    set_table_borders(table)
    for c, h in enumerate(headers):
        set_cell_text(table.cell(0, c), h, bold=True)
        set_cell_shading(table.cell(0, c), "D9EAF7")
    for r, row in enumerate(rows, 1):
        for c, text in enumerate(row):
            set_cell_text(table.cell(r, c), text)
            if c == 0:
                set_cell_shading(table.cell(r, c), "F2F6FA")
    doc.add_paragraph()


def add_boundary_table(doc):
    headers = ["边界项", "责任主体", "职责范围", "不承担/交由其他模块"]
    rows = [
        [
            "云接入客户端",
            "机器人端统一客户端",
            "拉取双光云台 RTSP 或本地媒体源；完成设备协议适配、本地编码/转封装、发布 LiveKit 视频/音频 Track；上报媒体状态；接收启动、停止、切换通道等指令。",
            "不负责平台用户权限决策；不负责前端观看 Token 分发；不直接管理平台侧媒体元数据索引。",
        ],
        [
            "Media Service",
            "平台媒体服务",
            "创建实时视频/语音会话；校验用户权限；生成机器人端发布 Token 和前端观看 Token；管理 Room、Participant、Track 元数据；处理超时、断流、释放和状态推送。",
            "不直接控制云台机械动作；不承载大媒体流传输；不把 RTSP 作为平台业务主协议。",
        ],
        [
            "媒体流通道",
            "LiveKit / WebRTC",
            "承载实时视频、实时音频、订阅、发布、SFU 转发、弱网适配和断连事件。",
            "不承载业务控制指令；不替代 MQTT/Robot Control Service 的设备控制职责。",
        ],
        [
            "控制指令通道",
            "MQTT / Robot Control Service",
            "承载开始推流、停止推流、切换通道、客户端回执、设备在线状态、异常上报等控制和事件。",
            "不传输视频或音频大媒体数据；不负责前端业务状态长连接展示。",
        ],
        [
            "业务状态通道",
            "WebSocket",
            "向 Web/App 推送会话创建、客户端 ACK、Track 发布、Streaming、Interrupted、Failed、Closed 等业务状态。",
            "不承载实时视频/音频流；不直接下发底层设备控制命令。",
        ],
        [
            "云台设备控制",
            "Robot Control Service",
            "负责云台方向、变焦、预置位、补光灯、雨刷、机械臂等设备动作控制。",
            "Media Service 只关心媒体通道 visible、thermal、fusion 等，不负责 PTZ 细节。",
        ],
        [
            "多路并发资源策略",
            "Media Service + LiveKit + 云接入客户端",
            "限制最大同时视频会话、高清订阅、发布 Track、对讲会话；同一机器人同一通道复用 Track；观看人数为 0 后延迟释放。",
            "不允许无约束地为每个观看者重复拉流或重复发布；不默认同时开启全部机器人双向对讲。",
        ],
        [
            "安全边界",
            "Auth Service + Media Service + 云接入客户端",
            "区分 API Token、机器人发布 Token、前端观看 Token；机器人只能发布自己的 Track；前端只能订阅授权机器人；Token 短时有效并可审计。",
            "不使用长期暴露的 RTSP 密码型地址作为平台主凭据；Room/Track 命名不暴露敏感信息。",
        ],
    ]
    table = doc.add_table(rows=1 + len(rows), cols=4)
    set_table_borders(table)
    for c, h in enumerate(headers):
        set_cell_text(table.cell(0, c), h, bold=True)
        set_cell_shading(table.cell(0, c), "D9EAF7")
    for r, row in enumerate(rows, 1):
        for c, text in enumerate(row):
            set_cell_text(table.cell(r, c), text)
            if c == 0:
                set_cell_shading(table.cell(r, c), "F2F6FA")
    doc.add_paragraph()


def add_media_asset_boundary_table(doc):
    headers = ["能力", "主路径", "兜底/扩展路径", "边界说明"]
    rows = [
        [
            "录制",
            "LiveKit Egress 录制平台实时媒体流，产物写入 MinIO，元数据进入检索库。",
            "云接入客户端本地录制后上传，适用于网络不稳定、平台 Egress 不可用或任务要求本地留存。",
            "录制由平台统一发起、鉴权、记录和审计，录制内容以平台实时媒体流为主。",
        ],
        [
            "抓拍",
            "服务端基于当前 LiveKit Track 截帧，保存操作员观看时发现异常的当前画面。",
            "云接入客户端基于源视频流截帧，适用于无人值守、计划任务或 LiveKit Track 不可用场景。",
            "实时观看异常抓拍归入实时视频派生动作；前端只发起抓拍请求，最终图片由服务端生成和入库。",
        ],
        [
            "回放",
            "播放平台已存储媒体文件，例如 Egress 录像、上传录像和抓拍图片。",
            "机器人本地未上传视频的临时访问或拉取归入远程点播模块。",
            "回放只处理平台已有媒体资产，不混入机器人本地未入库文件的直接播放职责。",
        ],
    ]
    table = doc.add_table(rows=1 + len(rows), cols=4)
    set_table_borders(table)
    for c, h in enumerate(headers):
        set_cell_text(table.cell(0, c), h, bold=True)
        set_cell_shading(table.cell(0, c), "D9EAF7")
    for r, row in enumerate(rows, 1):
        for c, text in enumerate(row):
            set_cell_text(table.cell(r, c), text)
            if c == 0:
                set_cell_shading(table.cell(r, c), "F2F6FA")
    doc.add_paragraph()


def build():
    doc = Document()
    style_document(doc)

    title = doc.add_paragraph(style="Title")
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.add_run("实时视频模块设计方案与方案选择说明")
    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = subtitle.add_run("具身智能装备集成管理平台 / 媒体服务模块")
    run.font.name = "Arial"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "微软雅黑")
    run.font.size = Pt(11)
    run.font.color.rgb = RGBColor(96, 96, 96)

    doc.add_heading("1. 背景与设计结论", level=1)
    doc.add_paragraph(
        "本方案面向具身智能装备集成管理平台中的媒体服务模块，重点讨论实时视频能力的接入方式、平台职责和后续语音对讲扩展。"
        "当前架构已经规划在机器人端统一部署云接入客户端，客户端承担协议转换、接入认证、通信保护、消息转发和网络流量管理等职责。"
        "机器人上装设备中的双光云台支持 RTSP 协议，因此需要在统一实时媒体体系和设备协议兼容之间取得平衡。"
    )
    add_kv_table(
        doc,
        [
            ("推荐结论", "选择方案 A 作为平台主方案：机器人端云接入客户端作为 LiveKit Publisher，统一发布视频/音频 Track。"),
            ("关键限定", "方案 A 不是要求双光云台或所有上装设备直接接入 LiveKit，而是由机器人端统一客户端完成 RTSP、厂商 SDK、本地设备等差异适配。"),
            ("方案 B 定位", "方案 B 保留为设备兼容和兜底能力，主要位于云接入客户端内部，必要时也可在平台侧提供旁路接入。"),
            ("一期落地", "双光云台 RTSP -> 云接入客户端 RTSP Adapter -> LiveKit Video Track -> 平台 Media Service/LiveKit SFU -> Web/App 观看。"),
        ],
    )

    doc.add_heading("2. 方案 A 与方案 B 的区别", level=1)
    add_comparison_table(doc)

    doc.add_heading("3. 为什么选择方案 A", level=1)
    add_number(doc, "与现有总体架构一致。架构图中已经规划云接入客户端、WebRTC(SRTP)、SFU、通信保护和接入认证，方案 A 正好让这些能力成为主链路。")
    add_number(doc, "更适合 N 个机器人接入。平台不需要理解每种机器人或上装设备的视频协议，只需要管理统一的 Room、Participant、Track、Token 和状态。")
    add_number(doc, "更适合实时视频与语音对讲一体化。后续双向语音对讲可以复用 LiveKit 房间、音频 Track、权限和会话控制模型。")
    add_number(doc, "更适合移动网络。机器人端主动出站连接平台，比平台反向访问机器人内网 RTSP 地址更容易穿透 4G/5G、NAT 和现场网络隔离。")
    add_number(doc, "降低平台媒体服务压力。平台媒体服务负责会话、鉴权、状态和调度，媒体转接由云接入客户端/边缘端承担，避免云侧集中维护大量拉流进程。")
    add_number(doc, "仍保留存量设备兼容。双光云台 RTSP 不被否定，而是作为客户端内部媒体源接入，经过适配后统一发布到 LiveKit。")

    doc.add_heading("4. 推荐总体架构", level=1)
    add_code_block(
        doc,
        [
            "双光云台 RTSP / 本地摄像头 / 麦克风 / 厂商 SDK",
            "        |",
            "机器人端云接入客户端",
            "  - 设备协议适配：RTSP、SDK、本地采集",
            "  - 媒体编码与转发：视频 Track、音频 Track",
            "  - 接入认证与通信保护",
            "  - 状态上报与指令接收",
            "        |",
            "LiveKit Publisher",
            "        |",
            "平台 Media Service + LiveKit SFU",
            "        |",
            "Web 管理控制台 / 手机 App / 业务应用",
        ],
    )

    doc.add_heading("5. 边界与职责划分", level=1)
    doc.add_paragraph(
        "在选择方案 A 后，实时视频设计的关键不只是媒体链路，还要明确平台、云接入客户端、控制服务和安全服务之间的边界。"
        "本节用于约束后续详细设计和联调范围，避免媒体服务承担设备控制、客户端适配或大规模转码中心等不应由其承担的职责。"
    )
    add_boundary_table(doc)

    doc.add_heading("5.1 媒体资产能力边界", level=2)
    doc.add_paragraph(
        "录制、抓拍、回放虽然都与实时视频相关，但主路径不同。"
        "本方案按业务触发场景划分责任：录制以平台实时媒体流为主，抓拍以当前观看画面为主，回放以平台已存储媒体文件为主。"
    )
    add_media_asset_boundary_table(doc)

    doc.add_heading("6. 实时视频模块设计", level=1)
    doc.add_heading("6.1 核心职责", level=2)
    add_bullet(doc, "平台前端：发起观看、展示视频画面、接收会话状态、执行停止观看和通道切换。")
    add_bullet(doc, "Media Service：创建视频会话、校验权限、生成 LiveKit Token、管理 Room/Track 元数据、推送 WebSocket 状态。")
    add_bullet(doc, "LiveKit：承载 SFU 转发、房间、Participant、Track、订阅和断连事件。")
    add_bullet(doc, "云接入客户端：连接双光云台 RTSP，发布 LiveKit 视频 Track，并接收平台控制指令。")
    add_bullet(doc, "Robot Control Service / MQTT：承载设备在线状态、媒体启动/停止指令、回执和异常上报。")

    doc.add_heading("6.2 会话状态", level=2)
    add_code_block(
        doc,
        [
            "INIT -> REQUESTING_CLIENT -> CLIENT_ACKED -> ROOM_READY -> STREAMING -> STOPPING -> CLOSED",
            "",
            "异常状态：",
            "INIT / REQUESTING_CLIENT / ROOM_READY -> TIMEOUT / FAILED",
            "STREAMING -> INTERRUPTED -> STREAMING 或 FAILED",
        ],
    )

    doc.add_heading("6.3 关键接口", level=2)
    add_kv_table(
        doc,
        [
            ("POST /api/media/video-sessions", "创建实时视频会话，参数包括 robotId、deviceId、channel、quality、reuse。"),
            ("GET /api/media/video-sessions/{sessionId}", "查询会话状态、Room、Track、观看人数、错误原因等。"),
            ("POST /api/media/video-sessions/{sessionId}/token", "获取观看 Token，前端凭 Token 加入 LiveKit Room。"),
            ("POST /api/media/video-sessions/{sessionId}/switch-channel", "切换双光云台通道，如 visible、thermal、fusion。"),
            ("POST /api/media/video-sessions/{sessionId}/stop", "停止实时视频会话，释放观看会话和客户端媒体资源。"),
        ],
    )

    doc.add_heading("6.4 双光云台通道模型", level=2)
    add_bullet(doc, "visible：可见光通道，适合常规巡检与人工观察。")
    add_bullet(doc, "thermal：热成像/红外通道，适合夜间、烟雾、温度异常等场景。")
    add_bullet(doc, "fusion：融合通道，如果云台或客户端具备融合能力，可作为扩展通道。")
    add_bullet(doc, "quality 建议支持 main、sub、auto。视频墙优先使用 sub，单路详情优先使用 main。")

    doc.add_heading("7. 与语音对讲的衔接", level=1)
    doc.add_paragraph(
        "选择方案 A 后，语音对讲不需要另起一套完全独立的实时通信体系。"
        "云接入客户端可以同时作为音频 Publisher/Subscriber，平台通过 LiveKit 房间和音频 Track 实现单机器人对讲、分组广播和少量多方协作。"
    )
    add_bullet(doc, "单机器人对讲：操作员与某台机器人建立双向音频。")
    add_bullet(doc, "分组广播：操作员向一组机器人下发语音，机器人端可按策略回传。")
    add_bullet(doc, "监听/抢占：默认视频墙静音，选中机器人后开启监听或主控对讲。")
    add_bullet(doc, "占用规则：同一机器人同一时间建议只允许一个主控对讲，其他用户可监听、排队或申请接管。")

    doc.add_heading("8. 方案 B 的保留方式", level=1)
    doc.add_paragraph(
        "方案 B 不作为平台第一阶段主链路，但仍然是必要的兼容能力。它应主要下沉到云接入客户端内部，作为不同媒体源的适配方式；"
        "在少数不具备客户端部署条件的机器人或固定摄像头场景中，平台侧也可以提供旁路 Media Adapter。"
    )
    add_bullet(doc, "客户端内部适配：RTSP、RTMP、GB28181、厂商 SDK、本地摄像头、本地麦克风。")
    add_bullet(doc, "平台侧旁路适配：用于无法部署云接入客户端的存量设备或临时接入设备。")
    add_bullet(doc, "对上保持统一：无论底层来源如何，平台都输出 video_session、audio_session、room、track、token、status。")

    doc.add_heading("9. 一期实施建议", level=1)
    add_number(doc, "先完成云接入客户端到 LiveKit 的最小发布链路，验证双光云台 RTSP 转 LiveKit Video Track。")
    add_number(doc, "Media Service 先实现会话创建、Token 分发、Room/Track 状态管理、WebSocket 状态推送。")
    add_number(doc, "先支持 visible 和 thermal 两个通道，fusion 作为扩展字段预留。")
    add_number(doc, "先支持多人观看复用同一路 Track，避免同一机器人重复拉流和重复发布。")
    add_number(doc, "先定义异常处理：机器人离线、客户端 ACK 超时、Track 中断、重复请求、观看人数为 0 后延迟释放。")
    add_number(doc, "语音对讲在实时视频稳定后复用同一房间/会话模型扩展，避免重复设计状态机。")

    doc.add_heading("10. 建议评审关注点", level=1)
    add_bullet(doc, "云接入客户端是否具备稳定运行 FFmpeg/GStreamer/LiveKit SDK 的资源条件。")
    add_bullet(doc, "机器人端 4G/5G 网络下的上行带宽、丢包、重连和码率自适应策略。")
    add_bullet(doc, "双光云台 RTSP 主码流/子码流地址获取方式、认证方式和通道切换方式。")
    add_bullet(doc, "LiveKit 部署规模、TURN/STUN、Nginx 代理、证书、端口和内外网访问策略。")
    add_bullet(doc, "录制、抓拍、回放是否从 LiveKit/Egress 获取，还是由客户端或平台侧另行拉流处理。")

    doc.add_heading("11. 最终结论", level=1)
    doc.add_paragraph(
        "在当前架构已经规划机器人端统一云接入客户端的前提下，实时视频模块应选择方案 A 作为主方案。"
        "该方案能够把平台实时媒体能力统一到 LiveKit/WebRTC 体系内，同时通过云接入客户端吸收 RTSP、厂商 SDK、本地设备等差异。"
        "因此，方案 A 负责平台标准化，方案 B 负责设备兼容；两者不是互斥关系，而是分层关系。"
    )

    doc.save(OUT)
    return OUT


if __name__ == "__main__":
    print(build())
